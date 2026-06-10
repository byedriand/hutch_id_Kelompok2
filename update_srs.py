from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Buka SRS yang sudah ada
doc = Document('srs_hutch.id_Kel2_REAL._1.3.docx')

print("✅ File SRS berhasil dibuka")
print(f"Total paragraphs saat ini: {len(doc.paragraphs)}")

# Analisis struktur
for i, para in enumerate(doc.paragraphs):
    if 'BAB' in para.text or '4.5' in para.text or 'BAB V' in para.text or 'Persyaratan Non-Fungsional' in para.text:
        print(f"Para {i}: {para.text[:80]}")
