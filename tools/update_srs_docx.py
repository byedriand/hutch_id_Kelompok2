from docx import Document

# Buka SRS
doc = Document('srs_hutch.id_Kel2_REAL._1.3.docx')

print("=" * 80)
print("🔄 UPDATE SRS v1.3 → v1.4 (JUNI 2026)")
print("=" * 80)

# 1. UPDATE VERSION DI HALAMAN JUDUL
print("\n1️⃣  Updating version on title page...")
version_updated = False
for para in doc.paragraphs:
    if "Versi 1.3" in para.text:
        for run in para.runs:
            if "1.3" in run.text:
                run.text = run.text.replace("1.3", "1.4")
                version_updated = True
            if "April 2026" in run.text:
                run.text = run.text.replace("April 2026", "Juni 2026")
        if version_updated:
            print(f"   ✅ Updated version to 1.4 (Juni 2026)")
            break

# 2. FIND POSITION UNTUK INSERT SECTION 4.6 
print("\n2️⃣  Finding insertion point for section 4.6...")
bab_v_idx = None

for i, para in enumerate(doc.paragraphs):
    if "BAB V" in para.text and "PERSYARATAN NON-FUNGSIONAL" in para.text:
        bab_v_idx = i
        print(f"   Found 'BAB V - PERSYARATAN NON-FUNGSIONAL' at index {i}")
        break

if bab_v_idx is None:
    print("   ❌ ERROR: Could not find BAB V!")
    exit(1)

# 3. INSERT SECTION 4.6 CONTENT SEBELUM BAB V
print("\n3️⃣  Inserting new section 4.6...")

# Ambil referensi ke paragraph BAB V
bab_v_para = doc.paragraphs[bab_v_idx]
bab_v_element = bab_v_para._element
parent = bab_v_element.getparent()

# Content yang akan ditambahkan
new_sections = [
    ("4.6 Manajemen Produk oleh Staff Penjualan", "Heading 1"),
    ("", "Normal"),
    ("4.6.1 Deskripsi dan Prioritas", "Heading 2"),
    ("Modul Manajemen Produk memungkinkan Staff Penjualan menambahkan produk baru ke katalog sistem. Produk yang ditambahkan akan tersedia untuk dipilih saat membuat Purchase Order baru dan ditampilkan dalam daftar produk untuk referensi.", "Normal"),
    ("Prioritas: SEDANG", "Normal"),
    ("", "Normal"),
    ("4.6.2 Alur Stimulus / Respons", "Heading 2"),
    ("Alur Utama:", "Normal"),
    ("1. Staff Penjualan membuka menu 'Tambah Produk' dari sidebar", "Normal"),
    ("2. Sistem menampilkan halaman dengan form dan grid daftar produk existing", "Normal"),
    ("3. Staff mengisi form: nama produk, harga jual, stok awal, keterangan, upload foto", "Normal"),
    ("4. Staff melihat preview foto secara real-time", "Normal"),
    ("5. Staff submit form dan sistem melakukan validasi data", "Normal"),
    ("6. Jika valid: sistem simpan produk, upload foto, generate notifikasi ke semua roles", "Normal"),
    ("7. Jika tidak valid: sistem tampilkan error message per field", "Normal"),
    ("", "Normal"),
    ("4.6.3 Spesifikasi Konten Form", "Heading 2"),
    ("Nama Produk: Text, Required, Unique, Max 100 karakter", "Normal"),
    ("Harga Jual: Currency format Rupiah (Rp), Required, harus > 0", "Normal"),
    ("Stok Awal: Integer, Required, nilai >= 0", "Normal"),
    ("Keterangan: Textarea, Optional, Max 500 karakter", "Normal"),
    ("Foto Produk: Format JPG/PNG, Max 5MB, dengan live preview instant", "Normal"),
    ("", "Normal"),
    ("4.6.4 Persyaratan Fungsional", "Heading 2"),
    ("REQ-FR-020: Input Produk Baru - Staff penjualan dapat membuka form tambah produk dengan validasi real-time", "Normal"),
    ("REQ-FR-021: Notifikasi Produk Baru - Sistem otomatis membuat notifikasi tipe 'produk_baru' ke semua user roles", "Normal"),
    ("REQ-FR-022: Grid Display Produk - Tampilkan semua produk dalam grid responsive dengan card design", "Normal"),
    ("REQ-FR-023: Upload Foto Produk - Support format JPG/PNG, max 5MB, validasi di client dan server", "Normal"),
    ("REQ-FR-024: Audit Logging - Setiap penambahan produk di-log di audit trail untuk compliance", "Normal"),
]

# Insert paragraphs sebelum BAB V
# Kita perlu insert paragraphs dengan memanipulasi XML langsung
inserted_count = 0
for text, style in new_sections:
    # Create new paragraph
    new_p = doc.add_paragraph(text, style=style)
    new_p_element = new_p._element
    
    # Remove dari akhir dokumen
    new_p_element.getparent().remove(new_p_element)
    
    # Insert sebelum BAB V
    parent.insert(parent.index(bab_v_element) + inserted_count, new_p_element)
    inserted_count += 1

print(f"   ✅ Inserted {inserted_count} new paragraphs before BAB V")

# 4. SAVE FILE  
output_file = 'srs_hutch.id_Kel2_REAL._1.4.docx'
doc.save(output_file)
print(f"\n✅ File successfully saved: {output_file}")

print("\n" + "=" * 80)
print("🎉 UPDATE SELESAI!")
print("=" * 80)
print(f"\nFile baru tersimpan: {output_file}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print("\n✅ Perubahan yang dilakukan:")
print("  • Version: 1.3 → 1.4")
print("  • Tanggal: April 2026 → Juni 2026")
print("  • Section baru: 4.6 Manajemen Produk oleh Staff Penjualan")
print("  • Requirements baru: REQ-FR-020 s/d REQ-FR-024")
print("\n📝 Dokumentasi update tersedia di: SRS_UPDATE_DOCUMENT_v1.4.md")
